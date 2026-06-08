#!/usr/bin/env python3
"""Convert project markdown plan to Word (.docx)."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_formatted_runs(paragraph, text: str, bold_default=False, italic_default=False):
    """Parse **bold**, *italic*, and [link](url) inline markup."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\)|`[^`]+`)"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            run = paragraph.add_run(label)
            run.italic = True
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if bold_default and paragraph.runs:
        paragraph.runs[0].bold = True
    if italic_default and paragraph.runs:
        for run in paragraph.runs:
            run.italic = True


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], lang: str):
    if lang == "mermaid":
        p = doc.add_paragraph()
        run = p.add_run(
            "[Diagram — dostępny w wersji Markdown; szczegóły w sekcji 4 harmonogramu.]"
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
        return

    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)
    doc.add_paragraph()


def convert_md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    i = 0
    table_buffer: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if not is_table_separator(stripped):
                table_buffer.append(parse_table_row(stripped))
            i += 1
            continue
        elif table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Zapisano: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "plan-pracy-z-klientem.md"
    out = root / "plan-pracy-z-klientem.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    convert_md_to_docx(md, out)
